from django.shortcuts import render, redirect, get_object_or_404
from django.http import HttpResponse
from django.contrib import messages
from .models import Student, StudentCharacteristic, Relative
from .forms import StudentForm, RelativeForm, CharacteristicUploadForm
from docx import Document
from django.core.paginator import Paginator
import os, io, csv
from django.conf import settings
from docx.shared import Inches
from django.db.models import Q
from datetime import datetime
from django.contrib.auth import authenticate, login
from django.contrib.auth.decorators import login_required

@login_required
def student_list(request):
    students = Student.objects.all().order_by('last_name', 'first_name')    
    search_query = request.GET.get('q', '').strip()
    course_filter = request.GET.get('course', '').strip()
    faculty_filter = request.GET.get('faculty', '').strip()

    if search_query:
        students = students.filter(
            Q(first_name__icontains=search_query) |
            Q(last_name__icontains=search_query) |
            Q(middle_name__icontains=search_query) |
            Q(student_id__icontains=search_query)
        )
    
    if course_filter:
        students = students.filter(course=course_filter)
    
    if faculty_filter:
        students = students.filter(faculty=faculty_filter)

    paginator = Paginator(students, 10)
    page_number = request.GET.get('page')
    page_obj = paginator.get_page(page_number)

    response = render(request, 'students/student_list.html', {
        'page_obj': page_obj,
        'course_choices': Student.COURSE_CHOICES,
        'faculty_choices': Student.FACULTY_CHOICES,
        'query': search_query,
        'selected_course': course_filter,
        'selected_faculty': faculty_filter,
    })

    list(messages.get_messages(request))
    return response

def login_view(request):
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            messages.success(request, 'Siz ulgama ustunlikli girdiniz!')
            response = redirect('student_list')
        else:
            messages.error(request, 'Ýalňyş ulanyjy ady ýa-da parol')
            response = render(request, 'students/login.html')
    else:
        response = render(request, 'students/login.html')
    
    list(messages.get_messages(request))
    return response

@login_required
def student_detail(request, pk):
    student = get_object_or_404(Student, pk=pk)
    relatives = student.relatives.all()
    response = render(request, 'students/student_detail.html', {
        'student': student,
        'relatives': relatives
    })
    
    list(messages.get_messages(request))
    return response

@login_required
def save_characteristics_doc(student):
    document = Document()
    document.add_heading(f'{student.last_name} {student.first_name}-yn hasiyetnamasy', 0)

    if student.photo:
        photo_path = os.path.join(settings.MEDIA_ROOT, student.photo.name)
        if os.path.exists(photo_path):
            try:
                document.add_picture(photo_path, width=Inches(2.0))
            except Exception:
                document.add_paragraph("Suraty goşmakda ýalňyşlyk ýüze çykdy.")
        else:
            document.add_paragraph("Surat tapylmady.")
    else:
        document.add_paragraph("Surat ýok.")

    document.add_paragraph(f"Talyp: {student.last_name} {student.first_name} {student.middle_name or ''}")
    document.add_paragraph(f"Kurs: {student.get_course_display()}")
    document.add_paragraph(f"Fakultet: {student.get_faculty_display()}")
    document.add_paragraph(f"Talyp ID: {student.student_id}")
    document.add_paragraph(f"Doglan senesi: {student.birth_date}")
    document.add_paragraph(f"Doglan ýeri: {student.birth_place or '—'}")
    document.add_paragraph(f"Kabul edilen ýyly: {student.admission_year or '—'}")
    document.add_paragraph(f"Häsiýetnama: {student.characteristics or '—'}")
    document.add_paragraph(f"Bellikler: {student.notes or '—'}")

    folder_path = os.path.join(settings.MEDIA_ROOT, 'hasiyetnamalar')
    os.makedirs(folder_path, exist_ok=True)

    file_path = os.path.join(folder_path, f'{student.last_name}_{student.first_name}_hasiyetnamasy_{datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')

    try:
        document.save(file_path)
        return file_path
    except Exception as e:
        raise Exception(f'Ýalňyşlyk: {str(e)}')

@login_required
def student_create(request):
    if request.method == 'POST':
        form = StudentForm(request.POST, request.FILES)
        if form.is_valid():
            student = form.save()
            try:
                file_path = save_characteristics_doc(student)
                messages.success(request, f'Talyp gosuldy we häsiýetnama {file_path} ýerinde saklandy!')
            except Exception as e:
                messages.error(request, f'Talyp gosuldy, ýöne häsiýetnama saklanmady: {str(e)}')
            response = redirect('student_list')
        else:
            response = render(request, 'students/student_form.html', {'form': form})
    else:
        form = StudentForm()
        response = render(request, 'students/student_form.html', {'form': form})
    
    list(messages.get_messages(request))
    return response

@login_required
def student_update(request, pk):
    student = get_object_or_404(Student, pk=pk)
    if request.method == 'POST':
        form = StudentForm(request.POST, request.FILES, instance=student)
        if form.is_valid():
            student = form.save()
            try:
                file_path = save_characteristics_doc(student)
                messages.success(request, f'Talyp redaktirlenýär we häsiýetnama {file_path} salgyda saklandy!')
            except Exception as e:
                messages.error(request, f'Talyp redaktirlenýär, ýöne häsiýetnama saklanmady: {str(e)}')
            response = redirect('student_list')
        else:
            response = render(request, 'students/student_form.html', {'form': form, 'student': student})
    else:
        form = StudentForm(instance=student)
        response = render(request, 'students/student_form.html', {'form': form, 'student': student})
    
    list(messages.get_messages(request))
    return response

@login_required
def student_delete(request, pk):
    student = get_object_or_404(Student, pk=pk)
    if request.method == 'POST':
        student.delete()
        messages.success(request, 'Talyp öçürildi!')
        response = redirect('student_list')
    else:
        response = render(request, 'students/student_confirm_delete.html', {'student': student})
    
    list(messages.get_messages(request))
    return response

@login_required
def generate_characteristics(request, pk):
    student = get_object_or_404(Student, pk=pk)    
    try:
        document = Document()
        
        document.add_heading(f'{student.last_name} {student.first_name}-yn hasiyetnamasy', 0)
        
        if student.photo:
            photo_path = os.path.join(settings.MEDIA_ROOT, student.photo.name)
            if os.path.exists(photo_path):
                try:
                    document.add_picture(photo_path, width=Inches(2.0))
                except Exception as e:
                    document.add_paragraph(f"Surat gosmakda yalnyslyk: {str(e)}")
            else:
                document.add_paragraph("Serwerde surat tapylmady")
        else:
            document.add_paragraph("Surat yok")
        
        document.add_paragraph(f"Talyp: {student.last_name} {student.first_name} {student.middle_name or ''}")
        document.add_paragraph(f"Kursy: {student.get_course_display()}")
        document.add_paragraph(f"Fakulteti: {student.get_faculty_display()}")
        document.add_paragraph(f"Talyp belgisi: {student.student_id}")
        document.add_paragraph(f"Doglan wagty: {student.birth_date}")
        document.add_paragraph(f"Doglan yeri: {student.birth_place or '—'}")
        document.add_paragraph(f"Kabul edilen yyly: {student.admission_year or '—'}")
        document.add_paragraph(f"Hasiyetnama: {student.characteristics or '—'}")
        document.add_paragraph(f"Bellik: {student.notes or '—'}")
        
        document.add_paragraph(f"Doredilen(uytgedilen) wagty: {datetime.now().strftime('%d.%m.%Y')}")
        
        buffer = io.BytesIO()
        document.save(buffer)
        buffer.seek(0)  
        
        response = HttpResponse(
            buffer.getvalue(),
            content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        )
        response['Content-Disposition'] = (
            f'attachment; '
            f'filename={student.last_name}_{student.first_name}_hasiyetnamasy_{datetime.now().strftime("%Y%m%d")}.docx'
        )
        
        buffer.close()
        
        return response
        
    except Exception as e:
        error_msg = f"Hasiyetnamany doretmekde yalnyslyk: {str(e)}"
        print(error_msg)  
        
        return HttpResponse(
            f"Hasiyetnamany doretmekde yalnyslyk: {str(e)}", 
            status=500
        )

@login_required
def upload_characteristic_select(request):
    students = Student.objects.all()
    response = render(request, 'students/upload_characteristic_select.html', {
        'students': students
    })

    list(messages.get_messages(request))
    return response

@login_required
def upload_characteristic(request, pk):
    student = get_object_or_404(Student, pk=pk)
    if request.method == 'POST':
        form = CharacteristicUploadForm(request.POST, request.FILES)
        if form.is_valid():
            StudentCharacteristic.objects.filter(student=student).delete()
            characteristic = form.save(commit=False)
            characteristic.student = student
            characteristic.save()
            messages.success(request, 'Häsiýetnama goşuldy!')
            response = redirect('view_characteristic', pk=student.pk)
        else:
            response = render(request, 'students/upload_characteristic.html', {
                'form': form,
                'student': student
            })
    else:
        form = CharacteristicUploadForm()
        response = render(request, 'students/upload_characteristic.html', {
            'form': form,
            'student': student
        })
    
    list(messages.get_messages(request))
    return response

@login_required
def view_characteristic(request, pk):
    student = get_object_or_404(Student, pk=pk)
    characteristic = StudentCharacteristic.objects.filter(student=student).first()
    response = render(request, 'students/view_characteristic.html', {
        'student': student,
        'characteristic': characteristic
    })
    
    list(messages.get_messages(request))
    return response

@login_required
def relative_list(request, student_pk):
    student = get_object_or_404(Student, pk=student_pk)
    relatives = student.relatives.all().order_by('relationship')
    relatives = sorted(relatives, key=lambda r: Relative.RELATIONSHIP_ORDER.get(r.relationship, 999))
    response = render(request, 'students/relative_list.html', {
        'student': student,
        'relatives': relatives
    })
    
    list(messages.get_messages(request))
    return response

@login_required
def relative_create(request, student_pk):
    student = get_object_or_404(Student, pk=student_pk)
    if request.method == 'POST':
        form = RelativeForm(request.POST)
        if form.is_valid():
            relative = form.save(commit=False)
            relative.student = student
            relative.save()
            messages.success(request, 'Garyndas gosuldy!')
            response = redirect('relative_list', student_pk=student.pk)
        else:
            response = render(request, 'students/relative_form.html', {
                'form': form,
                'student': student
            })
    else:
        form = RelativeForm()
        response = render(request, 'students/relative_form.html', {
            'form': form,
            'student': student
        })
    
    list(messages.get_messages(request))
    return response

@login_required
def relative_update(request, student_pk, pk):
    student = get_object_or_404(Student, pk=student_pk)
    relative = get_object_or_404(Relative, pk=pk, student=student)
    if request.method == 'POST':
        form = RelativeForm(request.POST, instance=relative)
        if form.is_valid():
            form.save()
            messages.success(request, 'Garyndas redaktirlendi!')
            response = redirect('relative_list', student_pk=student.pk)
        else:
            response = render(request, 'students/relative_form.html', {
                'form': form,
                'student': student,
                'relative': relative
            })
    else:
        form = RelativeForm(instance=relative)
        response = render(request, 'students/relative_form.html', {
            'form': form,
            'student': student,
            'relative': relative
        })
    
    list(messages.get_messages(request))
    return response

@login_required
def relative_delete(request, student_pk, pk):
    student = get_object_or_404(Student, pk=student_pk)
    relative = get_object_or_404(Relative, pk=pk, student=student)
    if request.method == 'POST':
        relative.delete()
        messages.success(request, 'Garyndas öçürildi!')
        response = redirect('relative_list', student_pk=student.pk)
    else:
        response = render(request, 'students/relative_confirm_delete.html', {
            'student': student,
            'relative': relative
        })
    
    list(messages.get_messages(request))
    return response

@login_required
def generate_relatives_doc(request, student_pk):
    student = get_object_or_404(Student, pk=student_pk)
    relatives = student.relatives.all().order_by('relationship')
    relatives = sorted(relatives, key=lambda r: Relative.RELATIONSHIP_ORDER.get(r.relationship, 999))

    document = Document()
    document.add_heading(f'{student.last_name} {student.first_name}-yakyn garyndaslary barada maglumat', 0)
    document.add_paragraph(f"Talyp: {student.last_name} {student.first_name} {student.middle_name or ''}")
    document.add_paragraph(f"Kurs: {student.get_course_display()}")
    document.add_paragraph(f"Fakultet: {student.get_faculty_display()}")
    document.add_paragraph(f"Talyp ID: {student.student_id}")

    if relatives:
        table = document.add_table(rows=1, cols=6)
        table.style = 'Table Grid'
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'FAA'
        hdr_cells[1].text = 'Garyndashlyk'
        hdr_cells[2].text = 'Doglan wagty'
        hdr_cells[3].text = 'Doglan ýeri'
        hdr_cells[4].text = 'Işleýän ýeri'
        hdr_cells[5].text = 'Adres'

        for relative in relatives:
            row_cells = table.add_row().cells
            row_cells[0].text = relative.full_name
            row_cells[1].text = relative.get_relationship_display()
            row_cells[2].text = str(relative.birth_date)
            row_cells[3].text = relative.birth_place or '—'
            row_cells[4].text = relative.work_place or '—'
            row_cells[5].text = relative.address or '—'
    else:
        document.add_paragraph("Garyndashlar yok")

    buffer = io.BytesIO()
    document.save(buffer)
    buffer.seek(0)

    response = HttpResponse(
        content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
        content=buffer.getvalue()
    )
    response['Content-Disposition'] = f'attachment; filename={student.last_name}_{student.first_name}_uc_arkasy.docx'

    list(messages.get_messages(request))
    return response

@login_required
def save_relatives_doc(request, pk):
    if request.method == 'POST':
        student = get_object_or_404(Student, pk=pk)
        folder_path = request.POST.get('folder_path', '').strip()

        safe_folder = os.path.join(settings.MEDIA_ROOT, 'docs', folder_path.replace('/', os.sep))
        if not safe_folder.startswith(os.path.join(settings.MEDIA_ROOT, 'docs')):
            messages.error(request, 'Ýalňyş ýol: papka docs/ içinde bolmaly.')
            response = redirect('relative_list', student_pk=pk)
        else:
            os.makedirs(safe_folder, exist_ok=True)

            doc = Document()
            doc.add_heading(f'{student.last_name} {student.first_name}-yakyn garyndaslary barada maglumat', 0)

            table = doc.add_table(rows=1, cols=6)
            table.style = 'Table Grid'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'FAA'
            hdr_cells[1].text = 'Garyndashlyk'
            hdr_cells[2].text = 'Doglan wagty'
            hdr_cells[3].text = 'Doglan ýeri'
            hdr_cells[4].text = 'Işleýän ýeri'
            hdr_cells[5].text = 'Adres'

            relatives = student.relatives.order_by('relationship')
            relatives = sorted(relatives, key=lambda r: Relative.RELATIONSHIP_ORDER.get(r.relationship, 999))
            for relative in relatives:
                row_cells = table.add_row().cells
                row_cells[0].text = relative.full_name
                row_cells[1].text = relative.get_relationship_display()
                row_cells[2].text = str(relative.birth_date)
                row_cells[3].text = relative.birth_place or '—'
                row_cells[4].text = relative.work_place or '—'
                row_cells[5].text = relative.address or '—'

            file_path = os.path.join(safe_folder, f'{student.last_name}_{student.first_name} uc arkasy {datetime.now().strftime("%Y%m%d_%H%M%S")}.docx')
            try:
                doc.save(file_path)
                messages.success(request, f'Dokument {file_path} -de saklandy.')
            except Exception as e:
                messages.error(request, f'Ýalňyşlyk: {str(e)}')
            
            response = redirect('relative_list', student_pk=pk)
                
        list(messages.get_messages(request))
        return response
    
    response = redirect('relative_list', student_pk=pk)
    
    list(messages.get_messages(request))
    return response

@login_required
def export_students_csv(request):
    students = Student.objects.all().order_by('last_name', 'first_name')
    search_query = request.GET.get('q', '').strip()
    course_filter = request.GET.get('course', '').strip()
    faculty_filter = request.GET.get('faculty', '').strip()

    filters = Q()
    if search_query:
        filters |= (Q(first_name__icontains=search_query) |
                    Q(last_name__icontains=search_query) |
                    Q(middle_name__icontains=search_query) |
                    Q(student_id__icontains=search_query))
    if course_filter:
        filters &= Q(course=course_filter)
    if faculty_filter:
        filters &= Q(faculty=faculty_filter)

    students = students.filter(filters)

    response = HttpResponse(content_type='text/csv')
    response['Content-Disposition'] = 'attachment; filename="talyplar.csv"'

    writer = csv.writer(response)
    writer.writerow(['Familiýasy', 'Ady', 'Atasynyň ady', 'Talyp ID', 'Kurs', 'Fakultet', 'Doglan senesi'])
    
    for student in students:
        writer.writerow([
            student.last_name,
            student.first_name,
            student.middle_name or '—',
            student.student_id,
            student.get_course_display(),
            student.get_faculty_display(),
            student.birth_date
        ])
    return response